"""
Layout-aware DOCX parser v2.

This module provides a comprehensive DOCX parser that preserves document structure,
formatting, and layout information in the IR v2 format.
"""

import io
import re
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.shared import Inches
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.run import Run as DocxRun
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement

from models.ir_v2 import (
    Document as IRDocument,
    DocumentMeta,
    Section,
    Header,
    Footer,
    Heading,
    Paragraph,
    ListItem,
    Table,
    Row,
    Cell,
    Figure,
    Textbox,
    Span,
    Anchor,
    TranslationStats,
    ParseResult,
    Block
)


def parse_docx(file_bytes: bytes) -> Tuple[IRDocument, str, Dict]:
    """
    Parse a DOCX file and return IR v2 document, language, and stats.
    
    Args:
        file_bytes: Raw DOCX file bytes
        
    Returns:
        Tuple of (Document, language, stats_dict)
    """
    with io.BytesIO(file_bytes) as fp:
        doc = Document(fp)
    
    # Initialize stats
    stats = TranslationStats()
    
    # Parse document metadata
    meta = _extract_metadata(doc)
    
    # Parse sections (for now, treat entire document as one section)
    sections = _parse_sections(doc, stats)
    
    # Create document
    ir_doc = IRDocument(meta=meta, sections=sections)
    
    # Detect language (simple heuristic for now)
    lang = _detect_language(ir_doc)
    
    # Convert stats to dict
    stats_dict = stats.dict()
    
    return ir_doc, lang, stats_dict


def _extract_metadata(doc: DocumentType) -> DocumentMeta:
    """Extract document metadata."""
    core_props = doc.core_properties
    
    # Count pages (rough estimate)
    page_count = len(doc.paragraphs) // 20  # Rough estimate
    
    # Count words
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    
    return DocumentMeta(
        title=core_props.title or None,
        author=core_props.author or None,
        created=core_props.created.isoformat() if core_props.created else None,
        modified=core_props.modified.isoformat() if core_props.modified else None,
        pages=max(1, page_count),
        word_count=word_count
    )


def _parse_sections(doc: DocumentType, stats: TranslationStats) -> List[Section]:
    """Parse document sections with simple, reliable text extraction."""
    blocks = []
    
    # Simple approach: process all paragraphs in order
    for para in doc.paragraphs:
        if para.text.strip():
            block = _parse_paragraph_simple(para)
            if block:
                blocks.append(block)
    
    # Process all tables
    for table in doc.tables:
        table_block = _parse_table_simple(table)
        if table_block:
            blocks.append(table_block)
    
    # Group blocks into logical sections based on headings
    sections = _group_into_sections_simple(blocks)
    
    # Update stats
    stats.paragraphs = len([b for b in blocks if isinstance(b, Paragraph)])
    stats.tables = len([b for b in blocks if isinstance(b, Table)])
    stats.sections = len(sections)

    if not sections:
        sections = [Section(index=0, blocks=[])]

    doc_sections = list(doc.sections)
    for idx, doc_section in enumerate(doc_sections):
        target_index = idx if idx < len(sections) else len(sections) - 1
        if target_index < 0:
            continue

        header_blocks = _parse_header_footer_part(doc_section.header)
        if header_blocks:
            sections[target_index].header = Header(blocks=header_blocks)
            stats.headers += 1

        footer_blocks = _parse_header_footer_part(doc_section.footer)
        if footer_blocks:
            sections[target_index].footer = Footer(blocks=footer_blocks)
            stats.footers += 1
    
    return sections


def _parse_paragraph_simple(para: DocxParagraph) -> Optional[Block]:
    """Parse a paragraph using simple, reliable heuristics."""
    text = para.text.strip()
    if not text:
        return None
    
    # Simple heading detection
    if _is_simple_heading_docx(text, para):
        level = _get_simple_heading_level_docx(text, para)
        spans = _extract_spans_simple(para)
        return Heading(level=level, spans=spans)
    
    # Simple list detection
    if _is_simple_list_item_docx(text):
        level = _get_simple_list_level_docx(text)
        ordered = _is_simple_ordered_list_docx(text)
        spans = _extract_spans_simple(para)
        return ListItem(level=level, ordered=ordered, spans=spans)
    
    # Regular paragraph
    spans = _extract_spans_simple(para)
    return Paragraph(spans=spans)


def _is_simple_heading_docx(text: str, para: DocxParagraph) -> bool:
    """Simple heading detection for DOCX."""
    # Check style first
    style_name = para.style.name.lower()
    if any(keyword in style_name for keyword in ['heading', 'title']):
        return True
    
    # Check for obvious heading patterns
    heading_patterns = [
        text.startswith(('Chapter', 'Section', 'Part', 'Appendix', 'CHAPTER', 'SECTION')),
        text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')),
        text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'IX.', 'X.')),
        text.startswith(('A.', 'B.', 'C.', 'D.', 'E.', 'F.', 'G.', 'H.', 'I.', 'J.')),
        text.startswith(('a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)', 'i)', 'j)')),
        text.startswith(('(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)', '(9)')),
        # Short lines that could be headings
        len(text) < 80 and not text.endswith('.') and not text.endswith(','),
        # All caps short lines
        text.isupper() and len(text) < 100,
    ]
    
    return any(heading_patterns)


def _get_simple_heading_level_docx(text: str, para: DocxParagraph) -> int:
    """Simple heading level detection for DOCX."""
    # Check style first
    style_name = para.style.name.lower()
    if 'heading 1' in style_name or 'title' in style_name:
        return 1
    elif 'heading 2' in style_name:
        return 2
    elif 'heading 3' in style_name:
        return 3
    elif 'heading 4' in style_name:
        return 4
    elif 'heading 5' in style_name:
        return 5
    elif 'heading 6' in style_name:
        return 6
    
    # Pattern-based detection
    if text.startswith(('Chapter', 'CHAPTER', 'Section', 'SECTION')):
        return 1
    elif text.startswith(('1.', '2.', '3.', '4.', '5.')):
        return 1
    elif text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
        return 2
    elif text.startswith(('A.', 'B.', 'C.', 'D.', 'E.')):
        return 3
    elif text.startswith(('a)', 'b)', 'c)', 'd)', 'e)')):
        return 4
    elif text.startswith(('(1)', '(2)', '(3)', '(4)', '(5)')):
        return 4
    else:
        # Default based on length
        if len(text) < 30:
            return 1
        elif len(text) < 50:
            return 2
        else:
            return 3


def _is_simple_list_item_docx(text: str) -> bool:
    """Simple list item detection for DOCX."""
    list_patterns = [
        text.startswith(('•', '-', '*', '◦', '▪', '▫')),
        text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')),
        text.startswith(('a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)', 'i)', 'j)')),
        text.startswith(('(a)', '(b)', '(c)', '(d)', '(e)', '(f)', '(g)', '(h)', '(i)', '(j)')),
        text.startswith(('(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)', '(9)')),
    ]
    
    return any(list_patterns)


def _get_simple_list_level_docx(text: str) -> int:
    """Simple list level detection for DOCX."""
    if text.startswith(('    ', '\t')):
        return 1
    elif text.startswith(('        ', '\t\t')):
        return 2
    else:
        return 0


def _is_simple_ordered_list_docx(text: str) -> bool:
    """Simple ordered list detection for DOCX."""
    return (text[0].isdigit() and '.' in text[:5]) or \
           text.startswith(('a)', 'b)', 'c)', 'i)', 'ii)', 'iii)'))


def _extract_spans_simple(para: DocxParagraph) -> List[Span]:
    """Extract spans with simple formatting detection."""
    spans = []
    
    for run in para.runs:
        text = _normalize_text(run.text)
        if not text:
            continue
            
        span = Span(
            text=text,
            bold=run.bold,
            italic=run.italic,
            underline=run.underline is not None
        )
        spans.append(span)
    
    # If no spans, create one with plain text
    if not spans:
        fallback_text = _normalize_text(para.text)
        if fallback_text:
            spans.append(Span(text=fallback_text))
    
    return spans


def _parse_table_simple(table: DocxTable) -> Optional[Table]:
    """Parse a table with simple structure detection."""
    if not table.rows:
        return None
    
    rows = []
    for docx_row in table.rows:
        cells = []
        for cell in docx_row.cells:
            # Simple cell content extraction
            cell_text = _normalize_text(cell.text)
            if cell_text:
                cell_blocks = [Paragraph(spans=[Span(text=cell_text)])]
            else:
                cell_blocks = [Paragraph(spans=[Span(text="")])]
            
            cells.append(Cell(blocks=cell_blocks))
        
        rows.append(Row(cells=cells))
    
    return Table(rows=rows)


def _group_into_sections_simple(blocks: List[Block]) -> List[Section]:
    """Group blocks into logical sections with simple heuristics."""
    sections = []
    current_section_blocks = []
    section_index = 0
    
    for block in blocks:
        # If we encounter a major heading, start a new section
        if isinstance(block, Heading) and block.level <= 2:
            # Save previous section if it has content
            if current_section_blocks:
                sections.append(Section(
                    index=section_index,
                    blocks=current_section_blocks
                ))
                section_index += 1
                current_section_blocks = []
        
        current_section_blocks.append(block)
    
    # Add the last section
    if current_section_blocks:
        sections.append(Section(
            index=section_index,
            blocks=current_section_blocks
        ))
    
    # If no sections were created, create one with all blocks
    if not sections:
        sections = [Section(index=0, blocks=blocks)]
    
    return sections


def _group_into_sections(blocks: List[Block]) -> List[Section]:
    """Group blocks into logical sections based on headings."""
    sections = []
    current_section_blocks = []
    section_index = 0
    
    for block in blocks:
        # If we encounter a major heading, start a new section
        if isinstance(block, Heading) and block.level <= 2:
            # Save previous section if it has content
            if current_section_blocks:
                sections.append(Section(
                    index=section_index,
                    blocks=current_section_blocks
                ))
                section_index += 1
                current_section_blocks = []
        
        current_section_blocks.append(block)
    
    # Add the last section
    if current_section_blocks:
        sections.append(Section(
            index=section_index,
            blocks=current_section_blocks
        ))
    
    # If no sections were created, create one with all blocks
    if not sections:
        sections = [Section(index=0, blocks=blocks)]
    
    return sections


def _parse_paragraph_improved(para: DocxParagraph, list_level: int = 0, in_list: bool = False) -> Optional[Block]:
    """Parse a paragraph with improved structure detection."""
    if not para.text.strip():
        return None
    
    # Check if it's a heading with improved detection
    heading_level = _get_heading_level_improved(para)
    if heading_level:
        spans = _extract_spans_improved(para)
        return Heading(level=heading_level, spans=spans)
    
    # Check if it's a list item with improved detection
    if _is_list_item_improved(para):
        spans = _extract_spans_improved(para)
        ordered = _is_ordered_list_improved(para)
        return ListItem(level=list_level, ordered=ordered, spans=spans)
    
    # Regular paragraph
    spans = _extract_spans_improved(para)
    return Paragraph(spans=spans)


def _get_heading_level_improved(para: DocxParagraph) -> Optional[int]:
    """Improved heading level detection."""
    style_name = para.style.name.lower()
    
    # Check for explicit heading styles
    if 'heading 1' in style_name or 'title' in style_name:
        return 1
    elif 'heading 2' in style_name:
        return 2
    elif 'heading 3' in style_name:
        return 3
    elif 'heading 4' in style_name:
        return 4
    elif 'heading 5' in style_name:
        return 5
    elif 'heading 6' in style_name:
        return 6
    
    # Check for implicit heading patterns
    text = para.text.strip()
    if text.startswith(('Chapter', 'Section', 'Part', 'Appendix')):
        return 1
    elif text.startswith(('1.', '2.', '3.', '4.', '5.')):
        return 1
    elif text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
        return 2
    elif text.startswith(('A.', 'B.', 'C.', 'D.', 'E.')):
        return 3
    
    # Check font size (if available)
    try:
        if para.runs:
            font_size = para.runs[0].font.size
            if font_size and font_size.pt > 16:
                return 1
            elif font_size and font_size.pt > 14:
                return 2
            elif font_size and font_size.pt > 12:
                return 3
    except:
        pass
    
    return None


def _is_list_item_improved(para: DocxParagraph) -> bool:
    """Improved list item detection."""
    # Check for list numbering or bullet points
    style_name = para.style.name.lower()
    
    # Check style-based lists
    if any(keyword in style_name for keyword in ['list', 'bullet', 'number']):
        return True
    
    # Check for list numbering in paragraph
    if para._element.getparent().tag.endswith('numPr'):
        return True
    
    # Check for common list patterns
    text = para.text.strip()
    list_patterns = [
        text.startswith(('•', '-', '*', '◦', '▪')),
        text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')),
        text.startswith(('a)', 'b)', 'c)', 'd)', 'e)', 'f)', 'g)', 'h)', 'i)', 'j)')),
        text.startswith(('(a)', '(b)', '(c)', '(d)', '(e)', '(f)', '(g)', '(h)', '(i)', '(j)')),
        text.startswith(('(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)', '(9)')),
    ]
    
    return any(list_patterns)


def _is_ordered_list_improved(para: DocxParagraph) -> bool:
    """Improved ordered list detection."""
    text = para.text.strip()
    
    # Check for numbered patterns
    if text[0].isdigit() and '.' in text[:5]:
        return True
    
    # Check for letter patterns
    if text.startswith(('a)', 'b)', 'c)', 'i)', 'ii)', 'iii)')):
        return True
    
    # Check for parenthesized patterns
    if text.startswith(('(1)', '(2)', '(3)', '(a)', '(b)', '(c)')):
        return True
    
    return False


def _extract_spans_improved(para: DocxParagraph) -> List[Span]:
    """Extract formatted spans with improved detection."""
    spans = []
    
    for run in para.runs:
        text = _normalize_text(run.text)
        if not text:
            continue
            
        span = Span(
            text=text,
            bold=run.bold,
            italic=run.italic,
            underline=run.underline is not None,
            link=_extract_link_improved(run)
        )
        spans.append(span)
    
    # If no spans, create one with plain text
    if not spans:
        fallback_text = _normalize_text(para.text)
        if fallback_text:
            spans.append(Span(text=fallback_text))
    
    return spans


def _extract_link_improved(run: DocxRun) -> Optional[str]:
    """Extract hyperlink with improved detection."""
    try:
        # Check if run has hyperlink
        for element in run._element:
            if element.tag.endswith('hyperlink'):
                r_id = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if r_id:
                    return f"#link_{r_id}"
    except:
        pass
    return None


def _parse_table_improved(table: DocxTable) -> Optional[Table]:
    """Parse a table with improved structure detection."""
    if not table.rows:
        return None
    
    rows = []
    for docx_row in table.rows:
        cells = []
        for cell in docx_row.cells:
            # Parse cell content with improved detection
            cell_blocks = []
            for para in cell.paragraphs:
                if para.text.strip():
                    # Determine if cell content is a heading, list, or paragraph
                    if _get_heading_level_improved(para):
                        level = _get_heading_level_improved(para)
                        spans = _extract_spans_improved(para)
                        cell_blocks.append(Heading(level=level, spans=spans))
                    elif _is_list_item_improved(para):
                        spans = _extract_spans_improved(para)
                        ordered = _is_ordered_list_improved(para)
                        cell_blocks.append(ListItem(level=0, ordered=ordered, spans=spans))
                    else:
                        spans = _extract_spans_improved(para)
                        cell_blocks.append(Paragraph(spans=spans))
            
            # If no blocks, create empty paragraph
            if not cell_blocks:
                cell_blocks.append(Paragraph(spans=[Span(text="")]))
            
            cells.append(Cell(blocks=cell_blocks))
        
        rows.append(Row(cells=cells))
    
    return Table(rows=rows)


def _parse_paragraph(para: DocxParagraph, list_level: int = 0, in_list: bool = False) -> Optional[Block]:
    """Parse a paragraph into appropriate block type."""
    if not para.text.strip():
        return None
    
    # Check if it's a heading
    heading_level = _get_heading_level(para)
    if heading_level:
        spans = _extract_spans(para)
        return Heading(level=heading_level, spans=spans)
    
    # Check if it's a list item
    if _is_list_item(para):
        spans = _extract_spans(para)
        ordered = _is_ordered_list(para)
        return ListItem(level=list_level, ordered=ordered, spans=spans)
    
    # Regular paragraph
    spans = _extract_spans(para)
    return Paragraph(spans=spans)


def _get_heading_level(para: DocxParagraph) -> Optional[int]:
    """Determine if paragraph is a heading and return level."""
    style_name = para.style.name.lower()
    
    if 'heading 1' in style_name:
        return 1
    elif 'heading 2' in style_name:
        return 2
    elif 'heading 3' in style_name:
        return 3
    elif 'heading 4' in style_name:
        return 4
    elif 'heading 5' in style_name:
        return 5
    elif 'heading 6' in style_name:
        return 6
    
    return None


def _is_list_item(para: DocxParagraph) -> bool:
    """Check if paragraph is a list item."""
    # Check for list numbering or bullet points
    style_name = para.style.name.lower()
    return ('list' in style_name or 
            para._element.getparent().tag.endswith('numPr') or
            any(run.text.startswith(('•', '-', '*', '1.', '2.', '3.')) for run in para.runs))


def _is_ordered_list(para: DocxParagraph) -> bool:
    """Check if list item is ordered (numbered)."""
    # Simple heuristic: check if starts with number
    text = para.text.strip()
    return text and text[0].isdigit()


def _extract_spans(para: DocxParagraph) -> List[Span]:
    """Extract formatted spans from paragraph."""
    spans = []
    
    for run in para.runs:
        text = _normalize_text(run.text)
        if not text:
            continue
            
        span = Span(
            text=text,
            bold=run.bold,
            italic=run.italic,
            underline=run.underline is not None,
            link=_extract_link(run)
        )
        spans.append(span)
    
    # If no spans, create one with plain text
    if not spans:
        fallback_text = _normalize_text(para.text)
        if fallback_text:
            spans.append(Span(text=fallback_text))
    
    return spans


def _extract_link(run: DocxRun) -> Optional[str]:
    """Extract hyperlink from run if present."""
    try:
        # Check if run has hyperlink
        for element in run._element:
            if element.tag.endswith('hyperlink'):
                r_id = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if r_id:
                    # In a full implementation, we'd resolve the relationship
                    return f"#link_{r_id}"
    except:
        pass
    return None


def _parse_table(table: DocxTable) -> Optional[Table]:
    """Parse a table into IR v2 format."""
    if not table.rows:
        return None
    
    rows = []
    for docx_row in table.rows:
        cells = []
        for cell in docx_row.cells:
            # Parse cell content
            cell_blocks = []
            for para in cell.paragraphs:
                if para.text.strip():
                    block = _parse_paragraph(para)
                    if block:
                        cell_blocks.append(block)
            
            # If no blocks, create empty paragraph
            if not cell_blocks:
                cell_blocks.append(Paragraph(spans=[Span(text="")]))
            
            cells.append(Cell(blocks=cell_blocks))
        
        rows.append(Row(cells=cells))
    
    return Table(rows=rows)


def _parse_header_footer_part(part) -> List[Block]:
    """Parse header or footer content and return IR blocks."""
    if part is None:
        return []

    blocks: List[Block] = []
    for para in part.paragraphs:
        block = _parse_paragraph(para)
        if block:
            blocks.append(block)
    return blocks


def _normalize_text(text: Optional[str]) -> str:
    """Normalize run text by removing line breaks that fragment paragraphs."""
    if not text:
        return ""
    cleaned = re.sub(r'[\r\n\v]+', ' ', text)
    cleaned = re.sub(r' {2,}', ' ', cleaned)
    return cleaned.strip()


def _detect_language(doc: IRDocument) -> str:
    """Detect document language (simple heuristic)."""
    # Extract all text
    all_text = []
    for section in doc.sections:
        for block in section.blocks:
            if isinstance(block, (Paragraph, Heading, ListItem)):
                for span in getattr(block, 'spans', []):
                    all_text.append(span.text)
    
    text_sample = ' '.join(all_text)[:1000]  # First 1000 chars
    
    # Simple heuristics
    if any(char in text_sample for char in 'ابپتثجچحخدذرزژسشصضطظعغفقکگلمنوهی'):
        return 'fa'  # Persian
    elif any(char in text_sample for char in 'абвгдеёжзийклмнопрстуфхцчшщъыьэюя'):
        return 'ru'  # Russian
    elif any(char in text_sample for char in 'àáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ'):
        return 'fr'  # French
    elif any(char in text_sample for char in 'äöüß'):
        return 'de'  # German
    else:
        return 'en'  # Default to English
