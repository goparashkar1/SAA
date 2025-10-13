"""
DOCX Writer v2 - Export IR v2 documents to DOCX format with layout preservation.

This module renders IR v2 documents back into Word while keeping structural
elements such as headers, footers, headings, lists, and tables intact.
"""

from typing import Any, Dict, Iterable, List, Optional
import io
import json
import logging
import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
import docx.oxml.shared

try:
    from docx.enum.section import WD_SECTION  # type: ignore
except ImportError:  # Older python-docx releases
    WD_SECTION = None

if WD_SECTION is None:
    CONTINUOUS = 0
    NEW_PAGE = 2
else:
    CONTINUOUS = WD_SECTION_START.CONTINUOUS
    NEW_PAGE = WD_SECTION_START.NEW_PAGE

from models.ir_v2 import (
    Document as IRDocument,
    Header,
    Footer,
    Heading,
    Paragraph,
    ListItem,
    Table,
    Cell,
    Figure,
    Textbox,
    Span,
    Block,
    Section,  # Added missing import
)
from translator_agent.utils.unitconv import mm_to_emu, pct_of, px_to_emu

logger = logging.getLogger(__name__)

DEFAULT_STYLE_MAP: Dict[str, str] = {
    "title": "Title",
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "paragraph": "Body",
    "quote": "Quote",
    "sidebar": "Sidebar",
    "list_item": "List Paragraph",
    "caption": "Caption",
    "table": "Table Grid",
}


def write_docx_from_layout_json(
    layout_json_path: Path,
    assets_dir: Optional[Path],
    out_path: Path,
    rtl: bool = False,
    style_map: Optional[Dict[str, str]] = None,
) -> None:
    """
    Render a DOCX document using Docling's layout JSON export to preserve the
    original document structure (columns, images, captions, tables, etc.).
    """
    if not layout_json_path.exists():
        raise FileNotFoundError(f"Layout JSON not found: {layout_json_path}")

    with layout_json_path.open(encoding="utf-8") as fp:
        layout_data = json.load(fp)

    combined_styles = dict(DEFAULT_STYLE_MAP)
    if style_map:
        combined_styles.update(style_map)

    doc = Document()
    _clear_document_body(doc)

    pages: Iterable[Dict[str, Any]] = layout_data.get("pages", [])
    pages = list(pages)
    if not pages:
        logger.warning("Layout JSON at %s contained no pages; writing empty document.", layout_json_path)
        _add_paragraph(doc, "No layout data available.", combined_styles.get("paragraph", "Body"), rtl)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return

    sections = doc.sections
    first_page = pages[0]
    first_section = sections[0]
    _apply_page_settings(first_section, first_page)
    _ensure_section_columns(first_section, int(first_page.get("columns", 1)))
    _render_page(doc, first_section, first_page, assets_dir, combined_styles, rtl)

    for page in pages[1:]:
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _apply_page_settings(section, page)
        _ensure_section_columns(section, int(page.get("columns", 1)))
        _render_page(doc, section, page, assets_dir, combined_styles, rtl)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _clear_document_body(doc: Document) -> None:
    body_element = doc._body._element
    for child in list(body_element):
        body_element.remove(child)


def _apply_page_settings(section, page: Dict[str, Any]) -> None:
    width_pt = page.get("width_pt")
    height_pt = page.get("height_pt")
    width_mm = page.get("width_mm")
    height_mm = page.get("height_mm")

    if width_pt:
        section.page_width = Pt(width_pt)
    elif width_mm:
        section.page_width = Emu(mm_to_emu(width_mm))
    if height_pt:
        section.page_height = Pt(height_pt)
    elif height_mm:
        section.page_height = Emu(mm_to_emu(height_mm))

    margins = page.get("margins_pt") or {}
    left = margins.get("left")
    right = margins.get("right")
    top = margins.get("top")
    bottom = margins.get("bottom")
    if left is not None:
        section.left_margin = Pt(left)
    if right is not None:
        section.right_margin = Pt(right)
    if top is not None:
        section.top_margin = Pt(top)
    if bottom is not None:
        section.bottom_margin = Pt(bottom)


def _ensure_section_columns(section, ncols: int) -> None:
    ncols = max(1, int(ncols) if ncols else 1)
    sect_pr = section._sectPr
    cols_elems = sect_pr.xpath("./w:cols")
    cols = cols_elems[0] if cols_elems else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(ncols))
    if ncols > 1:
        cols.set(qn("w:equalWidth"), "1")
    if not cols_elems:
        sect_pr.append(cols)


def _render_page(
    doc: Document,
    section,
    page: Dict[str, Any],
    assets_dir: Optional[Path],
    style_map: Dict[str, str],
    rtl: bool,
) -> None:
    blocks = page.get("blocks", []) or []
    header_blocks = [b for b in blocks if b.get("role") == "header"]
    footer_blocks = [b for b in blocks if b.get("role") == "footer"]
    body_blocks = [b for b in blocks if b.get("role") not in {"header", "footer"}]

    if header_blocks:
        _render_blocks(section.header, section, page, header_blocks, assets_dir, style_map, rtl)
    if footer_blocks:
        _render_blocks(section.footer, section, page, footer_blocks, assets_dir, style_map, rtl)
    _render_blocks(doc, section, page, body_blocks, assets_dir, style_map, rtl)


def _render_blocks(
    container,
    section,
    page: Dict[str, Any],
    blocks: Iterable[Dict[str, Any]],
    assets_dir: Optional[Path],
    style_map: Dict[str, str],
    rtl: bool,
) -> None:
    content_width = _content_width(section)
    for block in blocks:
        block_type = (block.get("type") or "paragraph").lower()

        if block_type in {"title", "h1", "h2", "h3", "h4", "h5", "h6", "paragraph", "quote", "sidebar", "caption"}:
            text = block.get("text", "")
            style_name = block.get("style") or style_map.get(block_type, style_map.get("paragraph", "Body"))
            paragraph = _add_paragraph(container, text, style_name, rtl)
            if block_type == "caption":
                paragraph.paragraph_format.keep_together = True
        elif block_type in {"list", "list_item"}:
            _render_list_block(container, block, style_map, rtl)
        elif block_type in {"image", "figure", "photo"}:
            _add_image(container, block, assets_dir, content_width, page, style_map, rtl)
        elif block_type == "table":
            _add_table(container, block, content_width, style_map, rtl)
        elif block_type == "line_break":
            container.add_paragraph("")
        else:
            text = block.get("text")
            if isinstance(text, str):
                style_name = block.get("style") or style_map.get("paragraph", "Body")
                _add_paragraph(container, text, style_name, rtl)


def _render_list_block(container, block: Dict[str, Any], style_map: Dict[str, str], rtl: bool) -> None:
    if block.get("type") == "list":
        items = block.get("items") or []
        for item in items:
            if isinstance(item, dict):
                _add_list_item(
                    container,
                    text=item.get("text", ""),
                    level=int(item.get("level", block.get("level", 0) or 0)),
                    rtl=rtl,
                    ordered=bool(item.get("ordered", block.get("ordered", False))),
                    style_map=style_map,
                )
            else:
                _add_list_item(
                    container,
                    text=str(item),
                    level=int(block.get("level", 0) or 0),
                    rtl=rtl,
                    ordered=bool(block.get("ordered", False)),
                    style_map=style_map,
                )
    else:
        _add_list_item(
            container,
            text=block.get("text", ""),
            level=int(block.get("level", 0) or 0),
            rtl=rtl,
            ordered=bool(block.get("ordered", False)),
            style_map=style_map,
        )


def _content_width(section) -> int:
    try:
        page_width = int(section.page_width)
        left = int(section.left_margin)
        right = int(section.right_margin)
        return max(page_width - left - right, 1)
    except Exception:  # pragma: no cover - defensive
        return int(section.page_width)


def _add_paragraph(container, text: str, style: Optional[str], rtl: bool):
    paragraph = container.add_paragraph(text or "")
    if style:
        try:
            paragraph.style = style
        except (KeyError, ValueError, AttributeError):
            logger.debug("Style %s not found; using default.", style)
    paragraph.paragraph_format.keep_together = True
    if rtl:
        paragraph.paragraph_format.right_to_left = True
    return paragraph


def _add_list_item(
    container,
    text: str,
    level: int,
    rtl: bool,
    ordered: bool,
    style_map: Dict[str, str],
) -> None:
    style_name = style_map.get("list_item", "List Paragraph")
    paragraph = container.add_paragraph()
    try:
        paragraph.style = "List Number" if ordered else ("List Bullet" if not ordered and style_name == "List Paragraph" else style_name)
    except (KeyError, ValueError, AttributeError):
        paragraph.style = style_name
    paragraph.add_run(text or "")

    if level:
        paragraph.paragraph_format.left_indent = Inches(0.25 * level)
    paragraph.paragraph_format.keep_together = True
    if rtl:
        paragraph.paragraph_format.right_to_left = True


def _add_image(
    container,
    block: Dict[str, Any],
    assets_dir: Optional[Path],
    content_width_emu: int,
    page: Dict[str, Any],
    style_map: Dict[str, str],
    rtl: bool,
) -> None:
    if not assets_dir or not assets_dir.exists():
        logger.warning("Skipping image block; assets directory unavailable: %s", assets_dir)
        return

    src = block.get("src") or block.get("path") or block.get("asset")
    if not src:
        logger.debug("Image block missing src: %s", block)
        return

    candidate = (assets_dir / src).resolve()
    if not candidate.exists():
        candidate = (assets_dir / Path(src).name).resolve()
    if not candidate.exists():
        logger.warning("Image asset %s not found for block %s", src, block)
        return

    width_emu = _resolve_image_width(block, content_width_emu, page)

    paragraph = container.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    if rtl:
        paragraph.paragraph_format.right_to_left = True

    run = paragraph.add_run()
    try:
        if width_emu:
            run.add_picture(str(candidate), width=Emu(width_emu))
        else:
            run.add_picture(str(candidate))
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning("Failed to add image %s (%s)", candidate, exc)
        paragraph.add_run(f"[Missing image: {src}]")
        return

    caption = block.get("caption") or block.get("title")
    if caption:
        caption_style = style_map.get("caption", DEFAULT_STYLE_MAP["caption"])
        caption_paragraph = _add_paragraph(container, caption, caption_style, rtl)
        caption_paragraph.paragraph_format.keep_together = True


def _resolve_image_width(block: Dict[str, Any], content_width_emu: int, page: Dict[str, Any]) -> Optional[int]:
    bbox = block.get("bbox") or {}
    width_px = block.get("width_px")
    width_pct = block.get("width_pct")

    if width_px:
        return px_to_emu(float(width_px))
    if width_pct:
        return pct_of(content_width_emu, float(width_pct))

    bbox_width = bbox.get("w") or bbox.get("width")
    if bbox_width:
        page_content_pt = _page_content_width_pt(page)
        if page_content_pt:
            ratio = float(bbox_width) / float(page_content_pt)
            ratio = max(0.05, min(ratio, 1.0))
            return pct_of(content_width_emu, ratio)
    return content_width_emu


def _page_content_width_pt(page: Dict[str, Any]) -> Optional[float]:
    width_pt = page.get("width_pt")
    if not width_pt and page.get("width_mm"):
        width_pt = float(page["width_mm"]) * 72.0 / 25.4
    margins = page.get("margins_pt") or {}
    if width_pt is None:
        return None
    left = margins.get("left", 72)
    right = margins.get("right", 72)
    return max(width_pt - left - right, 1)


def _add_table(
    container,
    block: Dict[str, Any],
    content_width_emu: int,
    style_map: Dict[str, str],
    rtl: bool,
) -> None:
    rows = block.get("rows") or block.get("data")
    if not rows:
        return
    rows = [list(row) for row in rows]
    ncols = len(rows[0])

    table = container.add_table(rows=len(rows), cols=ncols)
    table.style = style_map.get("table", DEFAULT_STYLE_MAP["table"])
    table.autofit = False

    header = block.get("header", True)
    widths_pct = block.get("column_widths_pct") or block.get("widths_pct") or block.get("col_pct")

    for r_idx, row in enumerate(rows):
        for c_idx, cell_value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = "" if cell_value is None else str(cell_value)
            for paragraph in cell.paragraphs:
                if rtl:
                    paragraph.paragraph_format.right_to_left = True
            if header and r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                _shade_cell(cell, "E6E6E6")

    if widths_pct:
        widths_iterable = widths_pct
    else:
        widths_iterable = [1 / ncols] * ncols

    for idx, pct_value in enumerate(widths_iterable):
        width_emu = pct_of(content_width_emu, float(pct_value))
        for cell in table.columns[idx].cells:
            cell.width = Emu(width_emu)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd_elements = tc_pr.findall(qn("w:shd"))
    for shd in shd_elements:
        tc_pr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _write_document_content(
    doc: Document,
    ir_doc: IRDocument,
    start_section_index: int = 0,
) -> int:
    section_index = start_section_index
    prev_width = ir_doc.meta.page_width
    prev_height = ir_doc.meta.page_height

    for section in ir_doc.sections:
        curr_width = section.page_width or prev_width
        curr_height = section.page_height or prev_height

        doc_section = _get_or_add_section(doc, section_index, section, prev_width, prev_height)
        doc_section.page_width = Inches(curr_width)
        doc_section.page_height = Inches(curr_height)

        if section.header:
            _populate_header(doc_section, section.header)

        if section.footer:
            _populate_footer(doc_section, section.footer)

        for block in section.blocks:
            _write_block(doc, block)

        prev_width = curr_width
        prev_height = curr_height
        section_index += 1

    return section_index


def _get_or_add_section(doc: Document, index: int, ir_section: Optional[Section], prev_width: float, prev_height: float):
    if index < len(doc.sections):
        return doc.sections[index]

    if ir_section is None:
        start_type = NEW_PAGE
    else:
        curr_width = ir_section.page_width or prev_width
        curr_height = ir_section.page_height or prev_height
        if math.isclose(curr_width, prev_width) and math.isclose(curr_height, prev_height):
            start_type = CONTINUOUS
        else:
            start_type = NEW_PAGE

    add_section = getattr(doc, "add_section", None)
    if callable(add_section):
        add_section(start_type)
    sections = doc.sections
    return sections[-1]


def _populate_header(doc_section, ir_header: Header) -> None:
    if ir_header.page_range == "first":
        doc_section.different_first_page_header_footer = True
        header_part = doc_section.first_page_header
    elif ir_header.page_range == "even":
        doc_section.even_and_odd_headers = True
        header_part = doc_section.even_page_header
    elif ir_header.page_range == "odd":
        doc_section.even_and_odd_headers = True
        header_part = doc_section.header
    else:  # "all" or None
        header_part = doc_section.header

    _reset_container(header_part)

    for block in ir_header.blocks:
        paragraph = header_part.add_paragraph()

        if isinstance(block, Heading):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, f"Heading {block.level}")
        elif isinstance(block, Paragraph):
            _apply_spans_to_paragraph(paragraph, block.spans)
        elif isinstance(block, ListItem):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, "List Number" if block.ordered else "List Bullet")
            if block.level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.25 * block.level)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)


def _populate_footer(doc_section, ir_footer: Footer) -> None:
    if ir_footer.page_range == "first":
        doc_section.different_first_page_header_footer = True
        footer_part = doc_section.first_page_footer
    elif ir_footer.page_range == "even":
        doc_section.even_and_odd_headers = True
        footer_part = doc_section.even_page_footer
    elif ir_footer.page_range == "odd":
        doc_section.even_and_odd_headers = True
        footer_part = doc_section.footer
    else:  # "all" or None
        footer_part = doc_section.footer

    _reset_container(footer_part)

    for block in ir_footer.blocks:
        paragraph = footer_part.add_paragraph()

        if isinstance(block, Heading):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, f"Heading {block.level}")
        elif isinstance(block, Paragraph):
            _apply_spans_to_paragraph(paragraph, block.spans)
        elif isinstance(block, ListItem):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, "List Number" if block.ordered else "List Bullet")
            if block.level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.25 * block.level)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)


def _reset_container(container) -> None:
    if container is None:
        return
    for paragraph in list(container.paragraphs):
        _delete_paragraph(paragraph)


def _delete_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    paragraph.text = ""


def _write_block(doc: Document, block: Block) -> None:
    if isinstance(block, Heading):
        _write_heading(doc, block)
    elif isinstance(block, Paragraph):
        _write_paragraph(doc, block)
    elif isinstance(block, ListItem):
        _write_list_item(doc, block)
    elif isinstance(block, Table):
        _write_table(doc, block)
    elif isinstance(block, Figure):
        _write_figure(doc, block)
    elif isinstance(block, Textbox):
        _write_textbox(doc, block)


def _write_heading(doc: Document, heading: Heading) -> None:
    paragraph = doc.add_heading(level=heading.level)
    _apply_spans_to_paragraph(paragraph, heading.spans)
    _set_style(paragraph, f"Heading {heading.level}")
    paragraph.paragraph_format.space_after = Pt(12)


def _write_paragraph(doc: Document, paragraph_model: Paragraph) -> None:
    paragraph = doc.add_paragraph()
    _apply_spans_to_paragraph(paragraph, paragraph_model.spans)
    _set_style(paragraph, "Normal")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    if paragraph_model.line_spacing:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = paragraph_model.line_spacing
    if paragraph_model.alignment:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = align_map.get(paragraph_model.alignment)


def _write_list_item(doc: Document, list_item: ListItem) -> None:
    paragraph = doc.add_paragraph()
    _apply_spans_to_paragraph(paragraph, list_item.spans)
    _set_style(paragraph, "List Number" if list_item.ordered else "List Bullet")
    if list_item.level > 0:
        paragraph.paragraph_format.left_indent = Inches(0.25 * list_item.level)
    paragraph.paragraph_format.space_after = Pt(4)


def _write_table(doc: Document, table_model: Table) -> None:
    if not table_model.rows:
        return

    column_counts = [sum(cell.colspan or 1 for cell in row.cells) for row in table_model.rows if row.cells]
    if not column_counts:
        return
    column_count = max(column_counts)
    docx_table = doc.add_table(rows=len(table_model.rows), cols=column_count)
    docx_table.style = "Table Grid"
    docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row in enumerate(table_model.rows):
        col_index = 0
        for cell in row.cells:
            target_cell = docx_table.cell(row_index, col_index)
            _write_cell_content(target_cell, cell)

            if cell.colspan and cell.colspan > 1:
                _merge_cells_horizontally(
                    docx_table, row_index, col_index, col_index + cell.colspan - 1
                )
            if cell.rowspan and cell.rowspan > 1:
                _merge_cells_vertically(
                    docx_table, row_index, col_index, row_index + cell.rowspan - 1
                )
            col_index += cell.colspan or 1

    docx_table.allow_autofit = True


def _write_cell_content(target_cell, cell_model: Cell) -> None:
    paragraphs = list(target_cell.paragraphs)
    if paragraphs:
        _clear_paragraph(paragraphs[0])
        base_paragraph = paragraphs[0]
    else:
        base_paragraph = target_cell.add_paragraph()

    for idx, block in enumerate(cell_model.blocks):
        paragraph = base_paragraph if idx == 0 else target_cell.add_paragraph()

        if isinstance(block, Paragraph):
            _apply_spans_to_paragraph(paragraph, block.spans)
        elif isinstance(block, Heading):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, f"Heading {block.level}")
        elif isinstance(block, ListItem):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, "List Number" if block.ordered else "List Bullet")
            if block.level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.25 * block.level)


def _write_figure(doc: Document, figure: Figure) -> None:
    if figure.image_data:
        stream = io.BytesIO(figure.image_data)
        width = Inches(figure.width or 4)
        height = Inches(figure.height or 3)
        doc.add_picture(stream, width=width, height=height)
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[Image: {figure.image_id}]")
        run.italic = True
        paragraph.paragraph_format.space_after = Pt(6)

    if figure.caption:
        caption = doc.add_paragraph()
        _apply_spans_to_paragraph(caption, figure.caption.spans)
        caption.style = "Caption"
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _write_textbox(doc: Document, textbox: Textbox) -> None:
    intro = doc.add_paragraph()
    _set_style(intro, "Intense Quote")
    intro.add_run("Textbox").bold = True

    for block in textbox.blocks:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)

        if isinstance(block, Paragraph):
            _apply_spans_to_paragraph(paragraph, block.spans)
        elif isinstance(block, Heading):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, f"Heading {block.level}")
        elif isinstance(block, ListItem):
            _apply_spans_to_paragraph(paragraph, block.spans)
            _set_style(paragraph, "List Number" if block.ordered else "List Bullet")
            if block.level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.25 * (block.level + 1))


def _apply_spans_to_paragraph(paragraph, spans: List[Span]) -> None:
    for span in spans:
        text = span.text or ""
        if not text:
            continue

        if span.link:
            run = _add_hyperlink(paragraph, span.link, text)
        else:
            run = paragraph.add_run(text)

        if span.bold:
            run.bold = True
        if span.italic:
            run.italic = True
        if span.underline:
            run.underline = True
        if span.font_size:
            run.font.size = Pt(span.font_size)
        if span.font_family:
            run.font.name = span.font_family
        if span.color and len(span.color) == 6:
            try:
                r = int(span.color[0:2], 16)
                g = int(span.color[2:4], 16)
                b = int(span.color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                pass


def _add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add blue color and underline by default for hyperlinks
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), "0000FF")
    rPr.append(c)

    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    # To apply additional styles, return the run
    return paragraph.runs[-1]


def _set_style(paragraph, style_name: str) -> None:
    if not style_name:
        return
    try:
        paragraph.style = style_name
    except (KeyError, ValueError, AttributeError):
        pass


def _write_no_api_placeholder(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Translation unavailable - export contains original content only.")
    run.italic = True


def _merge_cells_horizontally(table, row: int, start_col: int, end_col: int) -> None:
    start_cell = table.cell(row, start_col)
    for col in range(start_col + 1, end_col + 1):
        start_cell = start_cell.merge(table.cell(row, col))


def _merge_cells_vertically(table, start_row: int, col: int, end_row: int) -> None:
    start_cell = table.cell(start_row, col)
    for row in range(start_row + 1, end_row + 1):
        start_cell = start_cell.merge(table.cell(row, col))


def write_docx_from_html(
    html: str,
    out_path: Path,
    *,
    assets_dir: Optional[Path] = None,
    rtl: bool = True,
) -> None:
    """
    Convenience wrapper for rendering Markdown-derived HTML into DOCX using the
    lightweight HTML pipeline.
    """
    from .docx_writer import html_to_docx

    html_to_docx(
        html,
        out_path=out_path,
        rtl_override=rtl,
        assets_dir=assets_dir,
    )
