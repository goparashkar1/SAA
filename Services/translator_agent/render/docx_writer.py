# translator_agent/render/docx_writer.py
from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
import re

# Detect any RTL script char (Hebrew/Arabic/Persian ranges)
_RTL_RE = re.compile(r"[\u0590-\u05FF\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")

def _is_rtl(text: str) -> bool:
    return bool(_RTL_RE.search(text or ""))

def _apply_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), "1")
    pPr.append(bidi)

def _add_para(doc: Document, text: str, style: str | None = None, force_rtl: bool | None = None):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except Exception:
            pass  # style may not exist; ignore

    p.add_run(text)

    # Decide direction: explicit override > heuristic
    rtl = force_rtl if force_rtl is not None else _is_rtl(text)
    if rtl:
        _apply_rtl(p)
    # else: leave as default LTR

def _resolve_image_source(src: str | None, assets_dir: Optional[Path]) -> Optional[Path]:
    if not src:
        return None

    candidate = Path(src)
    if candidate.is_absolute() and candidate.exists():
        return candidate

    if assets_dir:
        cleaned = candidate
        if cleaned.is_absolute():
            cleaned = Path(cleaned.name)
        parts = list(cleaned.parts)
        if parts and parts[0] == assets_dir.name:
            cleaned = Path(*parts[1:]) if len(parts) > 1 else Path(cleaned.name)
        resolved = (assets_dir / cleaned).resolve()
        if resolved.exists():
            return resolved

    if candidate.exists():
        return candidate.resolve()
    return None


def _add_image(doc: Document, src: str | None, alt: str | None, assets_dir: Optional[Path]) -> None:
    path = _resolve_image_source(src, assets_dir)
    if not path:
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    try:
        run.add_picture(str(path))
    except Exception:
        # Invalid or unsupported image - skip gracefully
        paragraph._element.getparent().remove(paragraph._element)
        return

    if alt:
        caption = doc.add_paragraph(alt)
        caption.italic = True


def html_to_docx(
    html: str,
    out_path: str | Path | None = None,
    text_only: bool = False,
    *,
    rtl_override: bool | None = None,
    assets_dir: Optional[Path] = None,
) -> str:
    """
    Convert our simple HTML to a DOCX, applying RTL only to Farsi/Arabic/Hebrew paragraphs,
    and keeping English paragraphs LTR.
    """
    soup = BeautifulSoup(html, "lxml")
    doc = Document()

    # Set a readable default font; Word will still render English fine with this.
    normal = doc.styles["Normal"]
    if rtl_override:
        normal.font.name = "Vazirmatn"
    else:
        normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    body = soup.body or soup
    all_texts = []

    def _handle_paragraph(text: str) -> None:
        _add_para(doc, text, force_rtl=rtl_override)
        all_texts.append(text)

    for el in body.find_all(recursive=False):
        name = el.name
        if not name:
            continue

        if name in [f"h{i}" for i in range(1, 7)]:
            text = el.get_text()
            _add_para(doc, text, style=None, force_rtl=rtl_override)
            doc.paragraphs[-1].runs[0].bold = True
            if rtl_override is None and _is_rtl(text):
                _apply_rtl(doc.paragraphs[-1])
            all_texts.append(text)

        elif name == "p":
            images = list(el.find_all("img"))
            if images and not el.get_text(strip=True):
                for img in images:
                    _add_image(doc, img.get("src"), img.get("alt"), assets_dir)
                continue

            text = el.get_text()
            if text.strip():
                _handle_paragraph(text)

        elif name in ["ul", "ol"]:
            ordered = name == "ol"
            for li in el.find_all("li", recursive=False):
                li_images = list(li.find_all("img"))
                if li_images and not li.get_text(strip=True):
                    for img in li_images:
                        _add_image(doc, img.get("src"), img.get("alt"), assets_dir)
                    continue
                text = li.get_text()
                style = "List Number" if ordered else "List Bullet"
                _add_para(doc, text, style=style, force_rtl=rtl_override)
                all_texts.append(text)

        elif name == "blockquote":
            text = el.get_text()
            _handle_paragraph(text)

        elif name == "pre":
            text = el.get_text()
            _add_para(doc, text, force_rtl=False if rtl_override is None else rtl_override)
            all_texts.append(text)

        elif name == "img":
            _add_image(doc, el.get("src"), el.get("alt"), assets_dir)

        else:
            text = el.get_text()
            if text.strip():
                _handle_paragraph(text)

    plain_text = "\n".join(all_texts)

    if out_path is not None and not text_only:
        doc.save(str(out_path))

    return plain_text
