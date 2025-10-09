"""
IR v2 Models - Layout-aware intermediate representation for documents.
Enhanced with document formatting capabilities.
"""

from typing import List, Optional, Union, Literal, Any, Dict
from pydantic import BaseModel, Field, validator
from enum import Enum
import math
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
import io


class DocumentMeta(BaseModel):
    """Metadata about the document."""
    title: Optional[str] = None
    author: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    pages: int = 0
    word_count: int = 0
    page_width: float = 8.5  # inches
    page_height: float = 11.0  # inches


class Anchor(BaseModel):
    """Position anchor for elements that need placement information."""
    page: int = Field(ge=1, description="Page number (1-based)")
    x: float = Field(ge=0.0, le=1.0, description="Normalized x position (0-1)")
    y: float = Field(ge=0.0, le=1.0, description="Normalized y position (0-1)")
    width: Optional[float] = Field(default=None, ge=0.0, le=1.0, description="Normalized width (0-1)")
    height: Optional[float] = Field(default=None, ge=0.0, le=1.0, description="Normalized height (0-1)")

    def to_inches(self, page_width: float = 8.5, page_height: float = 11.0) -> Dict[str, float]:
        """Convert normalized coordinates to inches for Word document formatting."""
        return {
            'x': self.x * page_width,
            'y': self.y * page_height,
            'width': (self.width or 1.0) * page_width,
            'height': (self.height or 0.5) * page_height
        }


class Span(BaseModel):
    """Text span with formatting information."""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_size: Optional[float] = None  # points
    font_family: Optional[str] = None
    color: Optional[str] = None
    link: Optional[str] = None
    lang: Optional[str] = None


class Heading(BaseModel):
    """Heading block with level and formatted text."""
    level: int = Field(ge=1, le=6, description="Heading level (1-6)")
    spans: List[Span]
    anchor: Optional[Anchor] = None
    
    @property
    def text(self) -> str:
        return ''.join(span.text for span in self.spans)


class Paragraph(BaseModel):
    """Paragraph block with formatted text spans."""
    spans: List[Span]
    anchor: Optional[Anchor] = None
    alignment: Optional[Literal["left", "right", "center", "justify"]] = None
    line_spacing: Optional[float] = None
    
    @property
    def text(self) -> str:
        return ''.join(span.text for span in self.spans)


class ListItem(BaseModel):
    """List item with nesting level and formatting."""
    level: int = Field(ge=0, description="Nesting level (0-based)")
    ordered: bool = Field(description="True for numbered lists, False for bullet lists")
    spans: List[Span]
    anchor: Optional[Anchor] = None
    
    @property
    def text(self) -> str:
        return ''.join(span.text for span in self.spans)


class Cell(BaseModel):
    """Table cell containing blocks."""
    blocks: List[Union['Paragraph', 'ListItem', 'Heading']]  # Forward reference
    colspan: Optional[int] = Field(default=1, ge=1, description="Number of columns to span")
    rowspan: Optional[int] = Field(default=1, ge=1, description="Number of rows to span")
    dir: Optional[Literal["ltr", "rtl"]] = None


class Row(BaseModel):
    """Table row containing cells."""
    cells: List[Cell]


class Table(BaseModel):
    """Table with rows and cells."""
    rows: List[Row]
    anchor: Optional[Anchor] = None


class Figure(BaseModel):
    """Figure with image and optional caption."""
    image_id: str = Field(description="Unique identifier for the image")
    image_data: Optional[bytes] = None  # Actual image data
    image_format: Optional[str] = None  # 'png', 'jpg', etc.
    caption: Optional['Paragraph'] = None  # Forward reference
    anchor: Optional[Anchor] = None
    width: Optional[float] = None  # inches
    height: Optional[float] = None  # inches


class Textbox(BaseModel):
    """Text box with blocks and optional positioning."""
    blocks: List[Union['Paragraph', 'ListItem', 'Heading']]  # Forward reference
    anchor: Optional[Anchor] = None


class Header(BaseModel):
    """Document header with blocks."""
    blocks: List[Union['Paragraph', 'ListItem', 'Heading']]  # Forward reference
    page_range: Optional[str] = None  # "all", "first", "even", "odd"


class Footer(BaseModel):
    """Document footer with blocks."""
    blocks: List[Union['Paragraph', 'ListItem', 'Heading']]  # Forward reference
    page_range: Optional[str] = None  # "all", "first", "even", "odd"


# Union type for all possible block types (defined after all classes)
Block = Union[Heading, Paragraph, ListItem, Table, Figure, Textbox]


# Update forward references
Cell.update_forward_refs()
Figure.update_forward_refs()
Textbox.update_forward_refs()
Header.update_forward_refs()
Footer.update_forward_refs()


class Section(BaseModel):
    """Document section with optional header/footer and content blocks."""
    index: int = Field(ge=0, description="Section index (0-based)")
    header: Optional[Header] = None
    footer: Optional[Footer] = None
    blocks: List[Block]
    page_width: Optional[float] = 8.5
    page_height: Optional[float] = 11.0


class Document(BaseModel):
    """Root document model with metadata and sections."""
    meta: DocumentMeta
    sections: List[Section]

    def _add_spans_to_paragraph(self, spans: List[Span], p) -> None:
        for span in spans:
            run = p.add_run(span.text)
            if span.bold:
                run.bold = True
            if span.italic:
                run.italic = True
            if span.underline:
                run.underline = True
            if span.font_size:
                run.font.size = Pt(span.font_size)
            if span.font_family:
                run.font.name = span.font_family
            if span.color and len(span.color) == 6:
                try:
                    r = int(span.color[0:2], 16)
                    g = int(span.color[2:4], 16)
                    b = int(span.color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except ValueError:
                    pass  # Invalid color, skip
            # Link handling skipped for simplicity

    def _add_blocks_to_container(self, blocks: List[Block], container) -> None:
        for block in blocks:
            if isinstance(block, Heading):
                p = container.add_paragraph()
                p.style = f'Heading {block.level}'
                self._add_spans_to_paragraph(block.spans, p)
            elif isinstance(block, Paragraph):
                p = container.add_paragraph()
                self._add_spans_to_paragraph(block.spans, p)
                if block.alignment:
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    p.alignment = align_map.get(block.alignment)
                if block.line_spacing:
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p.paragraph_format.line_spacing = block.line_spacing
            elif isinstance(block, ListItem):
                style = 'List Number' if block.ordered else 'List Bullet'
                p = container.add_paragraph()
                p.style = style
                p.paragraph_format.left_indent = Inches(0.25 * (block.level + 1))
                p.clear()  # Clear default text if any
                self._add_spans_to_paragraph(block.spans, p)
            elif isinstance(block, Table):
                # Compute max columns considering colspans
                max_cols = max(sum(cell.colspan for cell in row.cells) for row in block.rows)
                table = container.add_table(rows=len(block.rows), cols=max_cols)
                for row_idx, row in enumerate(block.rows):
                    col_idx = 0
                    for cell in row.cells:
                        tc = table.cell(row_idx, col_idx)
                        if cell.colspan > 1 or cell.rowspan > 1:
                            tc.merge(table.cell(row_idx + cell.rowspan - 1, col_idx + cell.colspan - 1))
                        self._add_blocks_to_container(cell.blocks, tc)
                        col_idx += cell.colspan
            elif isinstance(block, Figure):
                if block.image_data:
                    stream = io.BytesIO(block.image_data)
                    p = container.add_paragraph()
                    run = p.add_run()
                    run.add_picture(stream, width=Inches(block.width or 4), height=Inches(block.height or 3))
                if block.caption:
                    self._add_blocks_to_container([block.caption], container)
            elif isinstance(block, Textbox):
                # Render inline for simplicity
                self._add_blocks_to_container(block.blocks, container)

    def to_docx(self, filename: str) -> None:
        doc = DocxDocument()
        # Set metadata
        if self.meta.title:
            doc.core_properties.title = self.meta.title
        if self.meta.author:
            doc.core_properties.author = self.meta.author
        # Set initial section properties
        section = doc.sections[0]
        section.page_width = Inches(self.meta.page_width)
        section.page_height = Inches(self.meta.page_height)
        current_section = section
        for sec in self.sections:
            if sec.index > 0:
                current_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
                current_section.page_width = Inches(sec.page_width or self.meta.page_width)
                current_section.page_height = Inches(sec.page_height or self.meta.page_height)
            # Handle header (simplified to use main header; extend for first/even/odd if needed)
            if sec.header:
                header = current_section.header
                if sec.header.page_range != "all":
                    if sec.header.page_range == "first":
                        current_section.different_first_page_headers = True
                        header = current_section.first_page_header
                    # Similarly for even/odd, but skipped for now
                self._add_blocks_to_container(sec.header.blocks, header)
            # Handle footer (similarly simplified)
            if sec.footer:
                footer = current_section.footer
                if sec.footer.page_range != "all":
                    if sec.footer.page_range == "first":
                        current_section.different_first_page_headers = True
                        footer = current_section.first_page_footer
                    # Similarly for even/odd
                self._add_blocks_to_container(sec.footer.blocks, footer)
            # Add content blocks
            self._add_blocks_to_container(sec.blocks, doc)
        doc.save(filename)


class GlossaryEntry(BaseModel):
    """Glossary entry for translation."""
    source: str
    target: str
    case_sensitive: bool = False
    exact: bool = False


class TranslationStats(BaseModel):
    """Statistics about the document parsing."""
    paragraphs: int = 0
    tables: int = 0
    figures: int = 0
    headers: int = 0
    footers: int = 0
    sections: int = 0
    word_count: int = 0


class ParseResult(BaseModel):
    """Result of parsing a document."""
    document: Document
    lang: str
    stats: TranslationStats