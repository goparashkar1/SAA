"""
Test IR v2 DOCX roundtrip functionality.

This test creates a sample DOCX with various elements, parses it with IR v2,
and then exports it back to DOCX to verify the roundtrip works correctly.
"""

import io
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from extract.docx_parser_v2 import parse_docx
from render.docx_writer_v2 import write_docx
from models.ir_v2 import Document as IRDocument, DocumentMeta, Section, Paragraph, Span, Heading, Table, Row, Cell


def create_test_docx() -> bytes:
    """Create a test DOCX with various elements."""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"
    
    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "Document Header"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page Footer"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add heading 1
    h1 = doc.add_heading("Main Title", level=1)
    
    # Add paragraph with formatting
    para = doc.add_paragraph()
    run1 = para.add_run("This is ")
    run2 = para.add_run("bold text")
    run2.bold = True
    run3 = para.add_run(" and this is ")
    run4 = para.add_run("italic text")
    run4.italic = True
    run5 = para.add_run(".")
    
    # Add heading 2
    h2 = doc.add_heading("Subsection", level=2)
    
    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Fill table
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Cell 1"
    table.cell(1, 1).text = "Cell 2"
    
    # Add another paragraph
    doc.add_paragraph("This is a regular paragraph with some content.")
    
    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def test_docx_parse():
    """Test DOCX parsing with IR v2."""
    # Create test DOCX
    docx_bytes = create_test_docx()
    
    # Parse with IR v2
    document, lang, stats = parse_docx(docx_bytes)
    
    # Assertions
    assert isinstance(document, IRDocument)
    assert document.meta.title == "Test Document"
    assert document.meta.author == "Test Author"
    assert len(document.sections) == 1
    
    section = document.sections[0]
    assert section.index == 0
    
    # Check for header
    assert section.header is not None
    assert len(section.header.blocks) > 0
    
    # Check for footer
    assert section.footer is not None
    assert len(section.footer.blocks) > 0
    
    # Check blocks
    blocks = section.blocks
    assert len(blocks) > 0
    
    # Find heading
    headings = [b for b in blocks if isinstance(b, Heading)]
    assert len(headings) > 0
    assert any(h.level == 1 for h in headings)
    
    # Find paragraphs
    paragraphs = [b for b in blocks if isinstance(b, Paragraph)]
    assert len(paragraphs) > 0
    
    # Find table
    tables = [b for b in blocks if isinstance(b, Table)]
    assert len(tables) > 0
    assert len(tables[0].rows) == 2
    
    # Check stats
    assert stats['paragraphs'] > 0
    assert stats['tables'] > 0
    assert stats['sections'] == 1
    
    print("✓ DOCX parsing test passed")
    return document


def test_docx_export():
    """Test DOCX export with IR v2."""
    # Create test document
    document, _, _ = parse_docx(create_test_docx())
    
    # Create placeholder target document
    target_meta = DocumentMeta(
        title="Translated Document",
        author="Test Author",
        pages=1,
        word_count=10
    )
    
    # Create placeholder content
    placeholder_para = Paragraph(spans=[Span(text="no credible API")])
    placeholder_section = Section(
        index=0,
        blocks=[placeholder_para]
    )
    
    target_document = IRDocument(
        meta=target_meta,
        sections=[placeholder_section]
    )
    
    # Export to DOCX
    docx_bytes = write_docx(document, target_document, "sequential")
    
    # Verify export
    assert len(docx_bytes) > 0
    
    # Parse the exported DOCX to verify content
    exported_doc = Document(io.BytesIO(docx_bytes))
    
    # Check that we have content
    paragraphs = exported_doc.paragraphs
    assert len(paragraphs) > 0
    
    # Check for "no credible API" text
    all_text = " ".join([p.text for p in paragraphs])
    assert "no credible API" in all_text
    
    print("✓ DOCX export test passed")
    return docx_bytes


def test_roundtrip():
    """Test complete roundtrip: DOCX -> IR v2 -> DOCX."""
    # Create original DOCX
    original_docx = create_test_docx()
    
    # Parse to IR v2
    document, lang, stats = parse_docx(original_docx)
    
    # Create target with placeholder
    target_meta = DocumentMeta(
        title="Translated Document",
        author="Test Author",
        pages=1,
        word_count=10
    )
    
    placeholder_para = Paragraph(spans=[Span(text="no credible API")])
    placeholder_section = Section(
        index=0,
        blocks=[placeholder_para]
    )
    
    target_document = IRDocument(
        meta=target_meta,
        sections=[placeholder_section]
    )
    
    # Export back to DOCX
    exported_docx = write_docx(document, target_document, "sequential")
    
    # Verify roundtrip
    assert len(exported_docx) > 0
    assert len(exported_docx) != len(original_docx)  # Should be different due to translation section
    
    # Parse exported DOCX to verify structure
    exported_doc = Document(io.BytesIO(exported_docx))
    paragraphs = exported_doc.paragraphs
    
    # Should have both original and translated content
    all_text = " ".join([p.text for p in paragraphs])
    assert "no credible API" in all_text
    
    print("✓ Roundtrip test passed")


def test_formatting_preservation():
    """Test that formatting is preserved in IR v2."""
    # Create DOCX with formatting
    doc = Document()
    
    # Add formatted paragraph
    para = doc.add_paragraph()
    run1 = para.add_run("Normal ")
    run2 = para.add_run("bold ")
    run2.bold = True
    run3 = para.add_run("and ")
    run4 = para.add_run("italic")
    run4.italic = True
    
    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    docx_bytes = bio.read()
    
    # Parse with IR v2
    document, _, _ = parse_docx(docx_bytes)
    
    # Find the paragraph
    section = document.sections[0]
    paragraphs = [b for b in section.blocks if isinstance(b, Paragraph)]
    assert len(paragraphs) > 0
    
    para = paragraphs[0]
    spans = para.spans
    
    # Check formatting
    assert len(spans) >= 4  # Should have multiple spans
    
    # Find bold and italic spans
    bold_spans = [s for s in spans if s.bold]
    italic_spans = [s for s in spans if s.italic]
    
    assert len(bold_spans) > 0
    assert len(italic_spans) > 0
    
    print("✓ Formatting preservation test passed")


if __name__ == "__main__":
    print("Running IR v2 DOCX roundtrip tests...")
    
    try:
        test_docx_parse()
        test_docx_export()
        test_roundtrip()
        test_formatting_preservation()
        
        print("\n🎉 All tests passed!")
        
    except Exception as e:
        print(f"\n❌ Test failed: {e}")
        raise
